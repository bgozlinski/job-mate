import io
import zipfile
from collections.abc import Sequence

import docx
import pytest

from app.services.extraction import (
    MAX_FILE_BYTES,
    ExtractionError,
    FileTooLarge,
    NoTextFound,
    UnsupportedFormat,
    extract_text,
)

RESUME = (
    "Backend engineer with eight years of experience building services in "
    "Python, PostgreSQL and Docker, working on payments and search platforms."
)
"""Long enough to clear MIN_EXTRACTED_CHARS, so a test that is not about the
threshold does not trip over it."""


def _pdf_with(text: str) -> bytes:
    """Build the smallest PDF that carries one line of extractable text.

    Written here rather than committed as a binary fixture: a reader can see
    what the test feeds the parser, and there is no file to keep in step with
    the code.
    """
    stream = b"BT /F1 12 Tf 72 720 Td (" + text.encode("ascii") + b") Tj ET"
    bodies = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length "
        + str(len(stream)).encode()
        + b">>stream\n"
        + stream
        + b"\nendstream",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []

    for number, body in enumerate(bodies, start=1):
        offsets.append(len(out))
        out += str(number).encode() + b" 0 obj" + body + b"endobj\n"

    started_at = len(out)
    out += b"xref\n0 " + str(len(bodies) + 1).encode() + b"\n0000000000 65535 f \n"
    out += b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets)
    out += b"trailer<</Size " + str(len(bodies) + 1).encode() + b"/Root 1 0 R>>\n"
    out += b"startxref\n" + str(started_at).encode() + b"\n%%EOF\n"

    return bytes(out)


def _docx_with(paragraphs: Sequence[str] = (), cells: Sequence[str] = ()) -> bytes:
    """Build a DOCX in memory, optionally putting some text in a table."""
    document = docx.Document()

    for text in paragraphs:
        document.add_paragraph(text)

    if cells:
        table = document.add_table(rows=len(cells), cols=1)
        for row, text in zip(table.rows, cells, strict=True):
            row.cells[0].text = text

    buffer = io.BytesIO()
    document.save(buffer)

    return buffer.getvalue()


def test_a_pdf_gives_back_its_text():
    assert RESUME in extract_text(_pdf_with(RESUME))


def test_a_docx_gives_back_its_paragraphs():
    assert RESUME in extract_text(_docx_with(paragraphs=[RESUME]))


def test_text_inside_a_docx_table_is_not_lost():
    """A two-column template is one table, and paragraphs alone would miss it."""
    extracted = extract_text(_docx_with(paragraphs=["Experience"], cells=[RESUME]))

    assert RESUME in extracted


def test_a_plain_text_file_is_read_as_it_is():
    assert extract_text(RESUME.encode("utf-8")) == RESUME


def test_a_byte_order_mark_does_not_reach_the_text():
    assert extract_text(RESUME.encode("utf-8-sig")) == RESUME


def test_a_windows_encoded_file_is_still_read():
    polish = f"{RESUME} Zażółć gęślą jaźń."

    assert extract_text(polish.encode("cp1250")) == polish


def test_the_result_is_normalized():
    assert extract_text(f"{RESUME}   \r\n\r\n".encode()) == RESUME


def test_the_old_word_format_is_refused_by_name():
    with pytest.raises(UnsupportedFormat, match="PDF or DOCX"):
        extract_text(b"\xd0\xcf\x11\xe0" + b"\x00" * 512)


def test_a_zip_that_is_not_a_docx_is_refused():
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("content.xml", RESUME)

    with pytest.raises(UnsupportedFormat):
        extract_text(buffer.getvalue())


def test_a_file_claiming_to_be_a_pdf_is_not_taken_at_its_word():
    with pytest.raises(UnsupportedFormat):
        extract_text(b"%PDF-1.4\nnot actually a pdf")


def test_binary_data_is_refused_rather_than_decoded():
    with pytest.raises(UnsupportedFormat):
        extract_text(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR")


def test_a_file_over_the_size_limit_is_refused_before_it_is_parsed():
    with pytest.raises(FileTooLarge):
        extract_text(b"a" * (MAX_FILE_BYTES + 1))


def test_a_pdf_with_too_many_pages_is_refused(monkeypatch):
    monkeypatch.setattr("app.services.extraction.MAX_PDF_PAGES", 0)

    with pytest.raises(FileTooLarge):
        extract_text(_pdf_with(RESUME))


def test_a_scan_is_reported_as_empty_rather_than_parsed():
    """A photograph in a PDF parses without raising and yields nothing."""
    with pytest.raises(NoTextFound):
        extract_text(_pdf_with("x"))


@pytest.mark.parametrize("data", [b"", b"   ", "krótko".encode()])
def test_too_little_text_is_never_a_document(data):
    with pytest.raises(NoTextFound):
        extract_text(data)


def test_every_refusal_shares_one_base_class():
    """The endpoint has one thing to catch, whatever went wrong."""
    for error in (UnsupportedFormat, FileTooLarge, NoTextFound):
        assert issubclass(error, ExtractionError)
