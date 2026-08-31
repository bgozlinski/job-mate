"""Turning an uploaded file into the text the rest of ingestion works on (FR-1).

Pure functions over bytes: no database, no HTTP and no filesystem, so every
format and every rejection can be tested from a literal.

This is the first of two layers. Getting characters out of a file is a
parser's job -- deterministic, free, and measured in milliseconds -- and only
a scan, where there are no characters to get, is worth a model. Reading the
structure out of those characters is the second layer and belongs to an LLM;
nothing here knows about it.

The split is also what keeps deduplication honest. A model asked to
transcribe the same PDF twice returns slightly different text, so a hash of
its output would make every re-upload a new document. The hash therefore
belongs on the uploaded bytes, which do not move when a parser is replaced.
"""

import io
import zipfile

import docx
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.services.chunking import normalize_content

MAX_FILE_BYTES = 5 * 1024 * 1024
"""What a resume can weigh. Generous for a document that is mostly text, and
low enough that a request cannot spend the container's memory. The check has
to be repeated where the upload is streamed: by the time bytes reach this
module they have already been read."""

MAX_PDF_PAGES = 50
"""A resume is not fifty pages. The limit is not about disk, it is about the
seconds a parser spends on a file whose only purpose is to be expensive."""

MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
"""A DOCX is a zip, so its declared size says nothing about what it expands
to. Sixteen kilobytes of nested nothing can decompress into gigabytes."""

MIN_EXTRACTED_CHARS = 100
"""Below this a parse is treated as having found nothing rather than as
having succeeded. A scanned resume -- a photograph wrapped in a PDF -- parses
without raising and yields a handful of stray characters, which would enter
the database as a document whose emptiness only surfaces later as a score of
zero. This is the hook the model fallback attaches to."""

PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"
OLE_MAGIC = b"\xd0\xcf\x11\xe0"
"""The compound-file header of the pre-2007 Office formats: .doc, .xls, .ppt."""

DOCX_ENTRY = "word/document.xml"
"""What separates a DOCX from every other zip, .odt and .pages included."""

TEXT_ENCODINGS = ("utf-8-sig", "cp1250")
"""Tried in order. utf-8-sig also reads plain UTF-8, and drops the byte order
mark an editor on Windows may have left in front of it. cp1250 is the second
guess for the same reason: a resume written in Polish on a Windows machine
and saved as plain text is not a hypothetical."""


class ExtractionError(Exception):
    """A file that cannot become text. Carries a message meant for the user."""


class UnsupportedFormat(ExtractionError):
    """The bytes are not a format this application reads."""


class FileTooLarge(ExtractionError):
    """The file is beyond one of the limits above."""


class NoTextFound(ExtractionError):
    """The format was read, and there was no text in it."""


def extract_text(data: bytes) -> str:
    """Return the normalised text of an uploaded file.

    There is deliberately no filename parameter. The format is decided by the
    first bytes of the file and nothing else: an extension is a claim made by
    whoever uploaded the file, and a .pdf that is really something else is the
    ordinary shape of an attack, not an edge case (NFR-1).

    The result is normalised through the same function ingestion hashes and
    splits with, so what this returns is already in the one form the rest of
    the pipeline agrees on.

    Raises ExtractionError -- one of UnsupportedFormat, FileTooLarge or
    NoTextFound -- for anything that cannot be turned into text. The three are
    separate because the answer to the user differs: change the format, send a
    smaller file, or send a document that is not a photograph.
    """
    if len(data) > MAX_FILE_BYTES:
        raise FileTooLarge(
            f"The file is larger than {MAX_FILE_BYTES // (1024 * 1024)} MB"
        )

    text = normalize_content(_extract(data))

    if len(text) < MIN_EXTRACTED_CHARS:
        raise NoTextFound(
            "No text could be read from the file. A scanned document has to be "
            "uploaded in a text form"
        )

    return text


def media_type(data: bytes) -> str:
    """Name the format the bytes actually are.

    Recorded next to the extracted text so the owner can be shown what they
    uploaded. Derived here rather than taken from the request: the
    Content-Type of a multipart part is set by the client, and a value the
    client chose is a claim, not a fact.

    Only called once extract_text has succeeded, so the format is already
    known to be one of these.
    """
    if data.startswith(PDF_MAGIC):
        return "application/pdf"

    if data.startswith(ZIP_MAGIC):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return "text/plain"


def _extract(data: bytes) -> str:
    """Dispatch on what the file actually is."""
    if data.startswith(PDF_MAGIC):
        return _from_pdf(data)

    if data.startswith(ZIP_MAGIC):
        return _from_docx(data)

    if data.startswith(OLE_MAGIC):
        raise UnsupportedFormat(
            "The old Word format is not supported. Save the document as PDF or "
            "DOCX and upload it again"
        )

    return _from_text(data)


def _from_pdf(data: bytes) -> str:
    """Read a PDF, page by page, in the order the file lists them."""
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages

        if len(pages) > MAX_PDF_PAGES:
            raise FileTooLarge(f"The PDF has more than {MAX_PDF_PAGES} pages")

        return "\n".join(page.extract_text() or "" for page in pages)
    except PdfReadError as exc:
        raise UnsupportedFormat("The file is not a readable PDF") from exc


def _from_docx(data: bytes) -> str:
    """Read a DOCX, including the tables a resume template hides text in.

    Paragraphs alone are not enough. Two-column resume templates are usually
    a single table, and a document read paragraph by paragraph gives back
    almost nothing -- which would surface as NoTextFound on a file that is
    full of text.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            if DOCX_ENTRY not in archive.namelist():
                raise UnsupportedFormat("The file is a zip archive, not a DOCX")

            if sum(entry.file_size for entry in archive.infolist()) > (
                MAX_UNCOMPRESSED_BYTES
            ):
                raise FileTooLarge("The DOCX expands to more than it should")
    except zipfile.BadZipFile as exc:
        raise UnsupportedFormat("The file is not a readable DOCX") from exc

    document = docx.Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return "\n".join(parts)


def _from_text(data: bytes) -> str:
    """Read a file that is already text, guessing only at its encoding."""
    if b"\x00" in data:
        raise UnsupportedFormat("The file is not a document this application reads")

    for encoding in TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise UnsupportedFormat("The text file is in an encoding that could not be read")
